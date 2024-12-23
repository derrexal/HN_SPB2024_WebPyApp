from io import BytesIO

from const import message_ok, mostly_correct, should_be_checked
from fuzzywuzzy import fuzz, process
from docx import Document
from transformers import pipeline
import pandas as pd
import json


def retrieve_delivery_time(pipe: pipeline, context: str) -> dict:
    return pipe(question='Какой срок поставки товара?', context=context)


def extract_table(file_bytes: bytes, table_num: int = 0):
    """Return Table from DOCX file format as DataFrame"""
    try:
        doc = Document(BytesIO(bytearray(file_bytes)))
    except Exception as e:
        return None

    table_count = len(doc.tables)
    if table_num >= table_count:
        raise BaseException('The count of tables is less than `table_num` argument')
    
    else:
        first_table = doc.tables[table_num]
        
        # Extract the table data into a list of lists
        table_data = []
        for row in first_table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        
        # Convert the list of lists into a Pandas DataFrame
        df = pd.DataFrame(table_data)
        
        # Optionally, set the first row as the header
        df.columns = df.iloc[0]
        df = df[1:]
        
        return df


def extract_tables(file_bytes: bytes) -> pd.DataFrame:
    """Return Table from DOCX file format as DataFrame"""
    doc = Document(BytesIO(bytearray(file_bytes)))
    table_count = len(doc.tables)

    if len(doc.tables) < 0:
        return None

    tables = []
    for i in range(table_count):
        first_table = doc.tables[i]

        # Extract the table data into a list of lists
        table_data = []
        for row in first_table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)

        # Convert the list of lists into a Pandas DataFrame
        df = pd.DataFrame(table_data)

        # Optionally, set the first row as the header
        df.columns = df.iloc[0]
        df = df[1:]

        tables.append(df)
    return tables


def compare_item_names(web_card: json, doc_df: pd.DataFrame) -> list:
    """
    Extract product items from web card and from document data
    And return the list of product items, that are not literally written in Document
    """
    # Extract item names from web card
    items = []
    for spec in web_card['specifications']:
        items.append(spec['title'])

    item_col = [i for i in doc_df.columns if 'Наименование' in i][0]
    doc_items = doc_df[item_col].to_list()
    
    missing_items = []
    for item in items:
        found = False
        for doc_item in doc_items:
            if item in doc_item:
                found = True
                print(f'[INFO] Item {item} PASSED THE CHECK (V)')
                break
        if not found:
            missing_items.append(item)
            print(f'[INFO] Item {item} HAS NOT PASSED THE CHECK (X)')
    return missing_items


def compare_item_quantity(web_card: json, doc_df: pd.DataFrame):
    # Extract item names from web card
    items = []
    for spec in web_card['specifications']:
        items.append(spec['title'])


# def check_if_text_in_docx(title: str, docx_text: str):
#     score = fuzz.partial_ratio(title.lower(), docx_text)
#     if score >= 90:
#         return message_ok, score
#     elif 90 > score >= 70:
#         return mostly_correct, score
#     else:
#         return should_be_checked, score


def find_product_name(product_item: str, data: pd.DataFrame, column: str='Наименование товара', thresh: int=90) -> tuple[str, int]:
    text, score = process.extractOne(product_item, data[column].to_list())
    if score >= thresh:
        df_product = data[data[column] == text]
        return df_product
    else:
        return "None"


def fuzzy_sim(str1: str, str2: str) -> int:
    """Возвращает скор по схожести"""
    return fuzz.partial_ratio(str1.lower(), str2.lower())


def get_verdict(x: int):
    if x >= 90:
        return message_ok
    elif 90 > x >= 70:
        return mostly_correct
    else:
        return should_be_checked


def extract_text_from_docx(docx_file: list[int] | bytes):
    try:
        doc = Document(BytesIO(bytearray(docx_file)))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        text = ' '.join(full_text)
        tables = extract_tables(docx_file)
        tables_str = ["\n".join(df_to_list_of_string(table)) for table in tables]

        return text + "\n".join(tables_str)
    except ValueError as ve:
        print(f"Ошибка при создании .docx: {ve}")
        return ""


def save_as_doc(byte_data: bytes) -> str:
    # Для старых .doc файлов не так просто извлечь текст, без использования дополнительных инструментов
    # Поэтому просто возвращаем байты как строку (пример: байты текстового документа)
    try:
        text = byte_data.decode('utf-8', errors='ignore')  # Предполагаем, что это текст в UTF-8
        print("Файл успешно сохранён как .doc")
        return text
    except Exception as e:
        print(f"Ошибка при сохранении в .doc: {e}")
        return ""


def df_to_list_of_string(data):
    return data.apply(lambda row: ' '.join(row.astype(str)), axis=1).to_list()
